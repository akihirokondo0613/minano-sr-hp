# 社内書式ページの生成器（リポジトリ内の複製）。正本の文面は data/shoshiki/forms.json。
# 元の生成器は 顧問先用書式ページ_作業/04_書式生成/render_forms.py（Desktop・Drive外）。
# ここでは HTML（編集用）だけを使う。docx/PDF の生成は元の作業フォルダで行う。
import datetime as _dt
import html as H
import json
import pathlib
import re

HERE = pathlib.Path(__file__).resolve().parent
REPO = HERE.parent.parent
DATA = json.loads(
    (REPO / "data" / "shoshiki" / "forms.json").read_text(encoding="utf-8")
)
META = DATA["meta"]
CO = META["company_placeholder"]

# ───────────────────────────────────────────────
# 記法のパース: [A|B] {d} {bNN} //
# ───────────────────────────────────────────────
TOK = re.compile(r"(\[[^\]]*\]|\{d\}|\{b\d+\}|\{h\d+\}|//)")


def parse(s):
    """→ [("text",str)|("choices",[..])|("date",)|("blank",mm)|("br",)]"""
    out = []
    for tok in TOK.split(s):
        if not tok:
            continue
        if tok.startswith("[") and tok.endswith("]"):
            out.append(("choices", tok[1:-1].split("|")))
        elif tok == "{d}":
            out.append(("date",))
        elif tok.startswith("{b"):
            out.append(("blank", int(tok[2:-1])))
        elif tok.startswith("{h"):
            out.append(("height", int(tok[2:-1])))  # 行の高さ(mm)。自由記述欄に使う
        elif tok == "//":
            out.append(("br",))
        else:
            out.append(("text", tok))
    return out


# ───────────────────────────────────────────────
# HTML
# ───────────────────────────────────────────────
CSS = """
:root{--ink:#1E2721;--ink2:#4A554D;--rule:#8A948C;--rule2:#C9D1CB;--shade:#EEF2EE;--brand:#1C5842;--accent:#2E9E63}
*{box-sizing:border-box}
html,body{margin:0;background:#E9ECE8}
body{font-family:"Hiragino Kaku Gothic ProN","Hiragino Sans","Yu Gothic","Noto Sans JP",sans-serif;color:var(--ink);font-size:10.5pt;line-height:1.6;line-break:strict;text-wrap:pretty}
.page{width:210mm;min-height:297mm;margin:12mm auto;background:#fff;padding:15mm 18mm 14mm;position:relative;box-shadow:0 2px 12px rgba(0,0,0,.12)}
h1{font-family:"Hiragino Mincho ProN","Yu Mincho",serif;font-size:16pt;font-weight:600;text-align:center;margin:0 0 4mm;letter-spacing:.06em}
.date{text-align:right;margin:0 0 3mm}
.to{margin:0 0 4mm}
.to div{margin:0}
.to .sender{display:table;margin:1.5mm 0 0 auto;line-height:1.7;text-align:left}  /* .to div{margin:0} より強くする */
/* 発信者欄（所在地・担当者・電話）は左揃えで、ラベルの直後から値が続く形。右端に寄せた箱の中で左揃えにする */
.sender .bl{text-align:left}
p{margin:0 0 1.6mm}
p.note{font-size:9pt;color:var(--ink2);margin-top:1.5mm}
table.f{width:100%;border-collapse:collapse;table-layout:fixed;margin:2.5mm 0 3mm}
table.f th,table.f td{border:0.6pt solid var(--rule);padding:1.6mm 2.2mm;vertical-align:middle;font-size:9.5pt;line-height:1.55}
table.f th{width:42mm;text-align:left;font-weight:700}
table.f th.w{width:52mm}
table.f td{min-height:8mm;height:8mm}
.opt{display:inline-block;white-space:nowrap}
.grp > :not(:last-child){margin-right:1.4em}
.opt .bx{display:inline-block;width:1.1em;cursor:pointer;user-select:none}
.lab{white-space:nowrap}
.bl{display:inline-block;border-bottom:0.6pt solid var(--ink);min-width:8mm;min-height:1.35em;line-height:1.35em;vertical-align:baseline;margin:0 .35em;padding:0 .3em .05em;text-align:center}
/* 入力前の空欄は箱の高さだけ確保し、下線を文字の基線より少し下に置く。入力後は文字の基線を周りの文字（ラベル）と揃える。
   以前は height 固定＋行送りより低い箱だったため、入力した文字が箱から下へはみ出して低く見えた */
.bl:empty{height:1.35em;vertical-align:-0.15em}
/* 会社情報（宛名の代表者・所在地・担当者・電話）が差し込まれた欄は下線を消す。空欄のままなら手書き用に下線を残す */
.bl.co-rep:not(:empty),.bl.co-addr:not(:empty),.bl.co-dept:not(:empty),.bl.co-tel:not(:empty){border-bottom-color:transparent}
.dt{white-space:nowrap}
.nb{white-space:nowrap}
.dt .bl{min-width:9mm;margin:0 .2em}
.dt .bl.y{min-width:12mm}
.checks{margin:1mm 0 3mm}
.checks h4{margin:0 0 1mm;font-size:10pt}
.checks div{padding-left:1.5em;text-indent:-1.5em;margin:0 0 .8mm}
.box{border:0.6pt solid var(--rule);margin:2.5mm 0 3mm}
.box .h{padding:1.4mm 2.2mm;font-size:9.5pt;font-weight:700;border-bottom:0.6pt solid var(--rule)}
.box .b{padding:2mm}
.t[contenteditable]:hover{outline:1px dashed var(--accent)}
.t[contenteditable]:focus{outline:2px solid var(--accent);background:#F3FBF6}
.bl[contenteditable]:focus{outline:2px solid var(--accent);background:#F3FBF6}
.bar{position:sticky;top:0;z-index:9;background:var(--brand);color:#fff;padding:8px 16px;font-size:12.5px;display:flex;gap:14px;align-items:center;flex-wrap:wrap}
.bar b{font-size:13.5px}
.bar button{font:inherit;padding:5px 12px;border:1px solid #fff;background:#fff;color:var(--brand);border-radius:3px;cursor:pointer;font-weight:700}
.bar span{opacity:.9}
.cfg{background:#fff;border-bottom:1px solid #C9D1CB;padding:12px 16px;display:none;font-size:13px}
.cfg.open{display:block}
.cfg .row{display:flex;flex-wrap:wrap;gap:8px 16px;align-items:center;max-width:980px;margin:0 auto}
.cfg label{display:flex;flex-direction:column;gap:2px;font-size:11.5px;color:#4A554D}
.cfg input{font:inherit;font-size:13.5px;padding:4px 8px;border:1px solid #C9D1CB;border-radius:3px;min-width:180px}
.cfg input.wide{min-width:320px}
.cfg .btns{display:flex;gap:8px;margin:8px auto 0;max-width:980px}
.cfg button{font:inherit;font-size:12.5px;padding:4px 10px;border:1px solid #1C5842;background:#fff;color:#1C5842;border-radius:3px;cursor:pointer}
.cfg .hint{font-size:11.5px;color:#4A554D;margin:6px auto 0;max-width:980px}
@media print{.cfg{display:none!important}}
@page{size:A4;margin:0}
@media print{html,body{background:#fff}.page{margin:0;box-shadow:none;width:210mm;height:297mm;page-break-after:always}.bar{display:none}.t[contenteditable]:hover,.t[contenteditable]:focus,.bl:focus{outline:none;background:transparent}}
"""

JS = """
(function(){
document.querySelectorAll('.bx').forEach(b=>b.addEventListener('click',()=>{b.textContent=b.textContent==='☐'?'☑':'☐';}));
window.pdf=function(){window.print();};
// ── 会社情報（ブラウザ内にだけ保存。どこにも送信しない） ──
// 関数は即時関数で包み、onclick から使うものだけ window に出す（他ページのJSと同名の const が衝突しないように）
const CO_KEY='shoshiki.company';
const CO_FIELDS=['name','title','rep','addr','tel','dept'];
function coLoad(){try{return JSON.parse(localStorage.getItem(CO_KEY)||'{}');}catch(e){return {};}}
function coSave(o){try{localStorage.setItem(CO_KEY,JSON.stringify(o));}catch(e){}}
function coApply(o){
  const set=(sel,v,def)=>document.querySelectorAll(sel).forEach(el=>{el.textContent=(v&&v.trim())?v.trim():(def||'');});
  set('.co-name',o.name,'【会社名】'); set('.co-title',o.title,'代表取締役'); set('.co-rep',o.rep,'');
  set('.co-addr',o.addr,''); set('.co-tel',o.tel,''); set('.co-dept',o.dept,'');
}
function coInit(){
  const o=coLoad(); coApply(o);
  CO_FIELDS.forEach(k=>{const i=document.getElementById('co-'+k); if(!i) return; i.value=o[k]||'';
    i.oninput=()=>{const c=coLoad(); c[k]=i.value; coSave(c); coApply(c);};});
}
window.coToggle=function(){document.getElementById('cfg').classList.toggle('open');};
window.coExport=function(){const o=coLoad(); const a=document.createElement('a');
  a.href='data:application/json;charset=utf-8,'+encodeURIComponent(JSON.stringify(o,null,1)); a.download='会社情報.json'; a.click();};
window.coImport=function(inp){const f=inp.files[0]; if(!f) return; const r=new FileReader();
  r.onload=()=>{try{const o=JSON.parse(r.result); coSave(o); coInit();}catch(e){alert('読み込めませんでした');}}; r.readAsText(f);};
window.coClear=function(){if(confirm('保存した会社情報を消しますか？')){localStorage.removeItem(CO_KEY); coInit();}};
coInit();
// 戻る／進むでキャッシュから復帰したときも保存内容を読み直す
window.addEventListener('pageshow',e=>{if(e.persisted) coInit();});
})();
"""

CFG_HTML = (
    '<div class="cfg" id="cfg"><div class="row">'
    '<label>会社名<input id="co-name" class="wide" placeholder="株式会社○○"></label>'
    '<label>代表者の役職<input id="co-title" placeholder="代表取締役"></label>'
    '<label>代表者氏名<input id="co-rep" placeholder="○○ ○○"></label>'
    '<label>所在地<input id="co-addr" class="wide" placeholder="富山市○○1-2-3"></label>'
    '<label>電話<input id="co-tel" placeholder="076-000-0000"></label>'
    '<label>担当部署・担当者<input id="co-dept" placeholder="総務部 ○○"></label>'
    '</div><div class="btns"><button onclick="coExport()">設定をファイルに書き出す</button>'
    '<button onclick="document.getElementById(\'co-file\').click()">設定ファイルを読み込む</button><input type="file" id="co-file" accept=".json" style="display:none" onchange="coImport(this)">'
    '<button onclick="coClear()">消去</button></div>'
    '<p class="hint">入力した会社情報はこのブラウザの中にだけ保存され、どこにも送信されません。書式の宛名・発信者欄に自動で入ります。別のPCで使うときは「書き出す」で保存したファイルを読み込んでください。</p></div>'
)


def esc(s):
    return H.escape(s)


def seg_html(seg, ed, lead=""):
    """lead: 選択肢群の直前の「ラベル：」。最初の選択肢と同じ行に束ねる"""
    kind = seg[0]
    ce = ' contenteditable="true"' if ed else ""
    if kind == "text":
        t = seg[1]
        if CO in t:
            # 「【会社名】（以下「会社」）と…」のように文中に含まれる場合も、会社名の部分だけ差し込み先にする
            out = []
            for i, part in enumerate(t.split(CO)):
                if i:
                    out.append(f'<span class="t co-name"{ce}>{esc(CO)}</span>')
                if part:
                    out.append(f'<span class="t"{ce}>{esc(part)}</span>')
            return "".join(out)
        return f'<span class="t"{ce}>{esc(t)}</span>'
    if kind == "choices":
        opts = [
            f'<span class="opt"><span class="bx">☐</span>{rich(o, ed)}</span>'
            for o in seg[1]
        ]
        if lead:
            opts[0] = (
                f'<span class="nb"><span class="t"{ce}>{esc(lead)}</span>{opts[0]}</span>'
            )
        return '<span class="grp">' + "".join(opts) + "</span>"
    if kind == "date":
        return f'<span class="dt"><span class="bl y"{ce}></span>年<span class="bl"{ce}></span>月<span class="bl"{ce}></span>日</span>'
    if kind == "blank":
        return f'<span class="bl" style="min-width:{seg[1]}mm"{ce}></span>'
    if kind == "br":
        return "<br>"
    return ""  # height は表側で処理


def rich(s, ed):
    segs = parse(s)
    out = []
    i = 0
    while i < len(segs):
        seg = segs[i]
        nxt = segs[i + 1] if i + 1 < len(segs) else None
        # 「生年月日：」＋日付、「氏名：」＋短い空欄 は同じ行に置く（末尾の短い語だけ nowrap で束ねる）
        if (
            seg[0] == "text"
            and nxt
            and nxt[0] == "choices"
            and seg[1].rstrip().endswith("：")
        ):
            txt = seg[1]
            m = re.search(r"[^\s　]{1,14}：$", txt.rstrip())
            head, tail = (txt[: m.start()], txt[m.start() :]) if m else ("", txt)
            if head:
                out.append(seg_html(("text", head), ed))
            out.append(seg_html(nxt, ed, lead=tail))
            i += 2
            continue
        if (
            seg[0] == "text"
            and nxt
            and (nxt[0] == "date" or (nxt[0] == "blank" and nxt[1] <= 30))
            and seg[1].rstrip().endswith(("：", "（"))
        ):
            txt = seg[1]
            m = re.search(r"[^\s　]{1,14}[：（]$", txt.rstrip())
            head, tail = (txt[: m.start()], txt[m.start() :]) if m else ("", txt)
            if head:
                out.append(seg_html(("text", head), ed))
            unit = ""
            nn = segs[i + 2] if i + 2 < len(segs) else None
            if nn and nn[0] == "text":
                m2 = re.match(r"([^　\s]{1,4})(?=$|　|\s)", nn[1])
                if m2 and not nn[1].startswith(("：", "（")):
                    unit = m2.group(1)
            out.append(
                '<span class="nb">'
                + seg_html(("text", tail), ed)
                + seg_html(nxt, ed)
                + (seg_html(("text", unit), ed) if unit else "")
                + "</span>"
            )
            if unit:
                rest = nn[1][len(unit) :]
                segs[i + 2] = ("text", rest)
            i += 2
            continue
        out.append(seg_html(seg, ed))
        i += 1
    return "".join(out)


def addressee_html(to, ed, f=None):
    ce = ' contenteditable="true"' if ed else ""
    if f is not None and "addr" in f:
        return (
            '<div class="to">'
            + "".join(f"<div>{rich(x, ed)}</div>" for x in f["addr"])
            + "</div>"
        )
    if to in ("会社", "保証人"):
        return f'<div class="to"><div class="t co-name"{ce}>{esc(CO)}</div><div><span class="t co-title"{ce}>代表取締役</span><span class="bl co-rep" style="min-width:36mm"{ce}></span>殿</div></div>'
    if to == "本人":
        return (
            f'<div class="to"><div><span class="bl" style="min-width:48mm"{ce}></span>殿</div>'
            f'<div class="sender"><div class="t co-name"{ce}>{esc(CO)}</div><div>所在地：<span class="bl co-addr" style="min-width:60mm"{ce}></span></div>'
            f'<div>担当者：<span class="bl co-dept" style="min-width:36mm"{ce}></span>　電話：<span class="bl co-tel" style="min-width:30mm"{ce}></span></div></div></div>'
        )
    if to == "主治医":
        return f'<div class="to"><div><span class="bl" style="min-width:44mm"{ce}></span>病院・医院</div><div><span class="bl" style="min-width:44mm"{ce}></span>先生　御机下</div></div>'
    if to == "会社←主治医":
        return f'<div class="to"><div class="t co-name"{ce}>{esc(CO)}</div><div>人事・労務担当者　殿</div></div>'
    if to == "社労士":
        return '<div class="to"><div>みなの社会保険労務士事務所　行</div></div>'
    return ""


def signature_rows(to, f=None):
    if f is not None and "sig" in f:
        return [tuple(r) for r in f["sig"]]
    if to == "会社":
        return [("所属", "{b70}"), ("氏名", "{b60}（署名）")]
    if to == "本人":
        return []  # 発信者欄は宛名の直後に置く
    if to == "双方":
        return [
            ("会社", CO + "　代表取締役{b40}（記名押印）"),
            ("従業員", "住所：{b50}//氏名：{b50}（署名）"),
        ]
    if to == "社労士":
        return [("会社名", ""), ("ご担当者（所属・氏名）", ""), ("電話・メール", "")]
    return []


def is_wide(f):
    """ラベルが12字以上なら列幅52mm（本文の表と署名欄の両方を数える）"""
    labels = [
        r["label"]
        for blk in f["blocks"]
        if blk["type"] == "fields"
        for r in blk["rows"]
    ]
    labels += [r[0] for r in signature_rows(f["to"], f)]
    return any(len(x) >= 12 for x in labels)


def fields_html(rows, ed, wide=False):
    h = ['<table class="f">']
    for r in rows:
        lab, val = (r["label"], r["value"]) if isinstance(r, dict) else r
        hm = re.search(r"\{h(\d+)\}", val)
        tall = (
            f' style="height:{hm.group(1)}mm"'
            if hm
            else (' style="height:13mm"' if val == "" else "")
        )
        val = re.sub(r"\{h\d+\}", "", val)
        h.append(
            f'<tr><th class="{"w" if wide else ""}"><span class="t"{' contenteditable="true"' if ed else ""}>{esc(lab)}</span></th>'
            + (
                f'<td class="t"{tall} contenteditable="true"></td>'
                if ed and val == ""
                else f"<td{tall}>{rich(val, ed)}</td>"
            )
            + "</tr>"
        )
    h.append("</table>")
    return "".join(h)


def form_html(f, ed):
    ce = ' contenteditable="true"' if ed else ""
    b = [f'<h1 class="t"{ce}>{esc(f["title"])}</h1>']
    b.append(
        f'<div class="date"><span class="dt"><span class="bl y"{ce}></span>年<span class="bl"{ce}></span>月<span class="bl"{ce}></span>日</span></div>'
    )
    b.append(addressee_html(f["to"], ed, f))
    for s in f["intro"]:
        b.append(f"<p>{rich(s, ed)}</p>")
    wide = is_wide(f)
    company = [blk for blk in f["blocks"] if blk.get("company")]
    for blk in f["blocks"]:
        if blk.get("company"):
            continue
        t = blk["type"]
        if t == "p":
            b.append(f"<p>{rich(blk['text'], ed)}</p>")
        elif t == "fields":
            b.append(fields_html(blk["rows"], ed, wide))
        elif t == "checks":
            b.append(
                f'<div class="checks"><h4 class="t"{ce}>{esc(blk["title"])}</h4>'
                + "".join(
                    f'<div><span class="bx">☐</span>　<span class="t"{ce}>{esc(i)}</span></div>'
                    for i in blk["items"]
                )
                + "</div>"
            )
        elif t == "box":
            b.append(
                f'<div class="box"><div class="h t"{ce}>{esc(blk["title"])}</div><div class="b t" style="min-height:{blk["height"]}mm"{ce}></div></div>'
            )
        elif t == "note":
            b.append(
                f'<p class="note">※ <span class="t"{ce}>{esc(blk["text"])}</span></p>'
            )
    sig = signature_rows(f["to"], f)
    if sig:
        b.append(fields_html(sig, ed, wide))
    for blk in company:
        b.append(fields_html(blk["rows"], ed, wide))
    b.append(
        ""  # 右下の社名（フッター）は廃止（2026-09-06・本人指示「場合によっては不自然」）
    )
    return '<div class="page">' + "".join(b) + "</div>"


def page(f, ed):
    bar = ""
    if ed:
        bar = (
            '<div class="bar"><b>編集モード</b><span>文字をクリックして打ち替え／☐をクリックで☑／下線の空欄に入力できます。</span>'
            '<button onclick="coToggle()">会社情報</button><button onclick="pdf()">印刷・PDFに保存</button><span>（印刷先を「PDFに保存」に）</span></div>'
            + CFG_HTML
        )
    return (
        f'<!doctype html>\n<html lang="ja"><head><meta charset="utf-8"><meta name="viewport" content="width=device-width,initial-scale=1">'
        f"<title>{esc(f['no'])} {esc(f['title'])}</title><style>{CSS}</style></head><body>{bar}{form_html(f, ed)}"
        f"<script>{JS}</script></body></html>\n"
    )


# ───────────────────────────────────────────────
# docx（選択肢は途中で折り返さないよう、長いときは1行ずつ）
# ───────────────────────────────────────────────
def build_docx(f, path):
    from docx import Document
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Mm, Pt

    FONT = "游ゴシック"

    def sf(run, size=10.5, bold=False):
        run.font.name = FONT
        run.font.size = Pt(size)
        run.font.bold = bold
        rpr = run._element.get_or_add_rPr()
        rf = rpr.find(qn("w:rFonts"))
        if rf is None:
            rf = OxmlElement("w:rFonts")
            rpr.append(rf)
        for k in ("w:eastAsia", "w:ascii", "w:hAnsi"):
            rf.set(qn(k), FONT)

    DATE_W = "＿＿＿＿年＿＿＿月＿＿＿日"

    def flat(s):
        """記法→Word用の文字列（本文10pt用）。空欄は下線文字、日付も下線、改行は \n、選択肢は □ 付き"""
        out = []
        for seg in parse(s):
            k = seg[0]
            if k == "text":
                out.append(seg[1])
            elif k == "date":
                out.append(DATE_W)
            elif k == "blank":
                out.append("＿" * max(2, int(seg[1] / 3.6 + 0.5)))
            elif k == "br":
                out.append("\n")
            elif k == "choices":
                out.append("　".join("□ " + flat(o) for o in seg[1]))
        return "".join(out)

    def cell_lines(s, limit):
        """セルに入れる行のリスト。全角スペースの位置だけで折り返し、
        「ラベル：＋空欄＋単位」「選択肢」は1語として途中で折らない。
        選択肢グループが1行に入るのに末尾だけ落ちるときは、直前の「ラベル：」ごとグループを次行へ移す"""
        BLANK = lambda mm: "＿" * max(2, int(mm / 3.2 + 0.5))  # noqa: E731  セル内は9pt・四捨五入

        def wlen(t):
            """表示幅（全角1・半角0.5）"""
            return sum(
                0.5 if ord(c) < 0x3000 or 0xFF61 <= ord(c) <= 0xFF9F else 1 for c in t
            )

        # 1) 語に分ける
        toks = []  # ("w", text) | ("sp",) | ("br",) | ("g", [options])
        for seg in parse(s):
            k = seg[0]
            if k == "br":
                toks.append(("br",))
            elif k == "choices":
                toks.append(("g", ["□ " + flat(o) for o in seg[1]]))
            elif k == "date":
                toks.append(("w", DATE_W))
            elif k == "blank":
                toks.append(("w", BLANK(seg[1])))
            elif k == "text":
                for part in re.split(r"(　)", seg[1]):
                    if part == "　":
                        toks.append(("sp",))
                    elif part:
                        toks.append(("w", part))
        # 2) 「ラベル：」＋空欄、空欄＋短い単位 をくっつける
        merged = []
        for t in toks:
            if t[0] == "w" and merged and merged[-1][0] == "w":
                prev = merged[-1][1]
                cur = t[1]
                joinable = (
                    prev.endswith(
                        ("：", "（", "→", "〜", "／", "〒", "－", "第", "金", "約")
                    )
                    or cur.startswith(
                        (
                            "）",
                            "円",
                            "名",
                            "日",
                            "時",
                            "分",
                            "回",
                            "か月",
                            "年",
                            "月",
                            "％",
                            "人",
                            "件",
                            "km",
                            "まで",
                            "頃",
                            "限り",
                            "を",
                            "に",
                            "から",
                            "へ",
                            "支部",
                        )
                    )
                    or (len(cur) <= 3 and not cur.endswith("："))
                    or cur.startswith("（")
                    or prev.endswith("＿")
                    or cur.startswith("＿")
                )
                if joinable:
                    merged[-1] = ("w", prev + cur)
                    continue
            merged.append(t)
        # 3) 幅で詰める
        lines = [""]

        def put(word):
            cur = lines[-1]
            sep = (
                ""
                if (
                    not cur
                    or cur.endswith(("：", "（"))
                    or word.startswith(("）", "」", "、", "。"))
                )
                else "　"
            )
            if cur.strip() and wlen(cur) + wlen(sep) + wlen(word) > limit:
                lines.append(word)
            else:
                lines[-1] = cur + sep + word

        for t in merged:
            if t[0] == "br":
                lines.append("")
            elif t[0] == "sp":
                continue
            elif t[0] == "w":
                put(t[1])
            elif t[0] == "g":
                opts = t[1]
                group = "　".join(opts)
                cur = lines[-1]
                m = re.search(r"([^　]{1,12}：)$", cur)
                label = m.group(1) if m else ""
                sep = "" if (not cur or cur.endswith(("：", "（"))) else "　"
                if wlen(cur) + wlen(sep) + wlen(group) <= limit:
                    lines[-1] = cur + sep + group
                elif wlen(label) + wlen(group) <= limit and cur.strip() != label:
                    lines[-1] = cur[: len(cur) - len(label)].rstrip("　")
                    if not lines[-1].strip():
                        lines.pop()
                    lines.append(label + group)
                else:
                    start = len(lines) - 1
                    for o in opts:
                        put(o)
                    # 最終行が選択肢1つだけなら、前の行の末尾の選択肢を送って2つ以上にする
                    if (
                        len(lines) - start >= 2
                        and lines[-1].count("□ ") == 1
                        and lines[-2].count("□ ") >= 2
                    ):
                        prev = lines[-2]
                        k = prev.rfind("□ ")
                        moved = prev[k:]
                        if wlen(moved) + 1 + wlen(lines[-1]) <= limit:
                            lines[-2] = prev[:k].rstrip("　")
                            lines[-1] = moved + "　" + lines[-1]
        return [ln for ln in lines if ln.strip()] or [""]

    def no_autospace(p):
        pPr = p._p.get_or_add_pPr()
        for tag in ("w:autoSpaceDE", "w:autoSpaceDN"):
            el = OxmlElement(tag)
            el.set(qn("w:val"), "0")
            pPr.append(el)

    def para(doc, text="", size=10, bold=False, align=None, after=3, ls=15):
        p = doc.add_paragraph()
        no_autospace(p)
        p.paragraph_format.line_spacing = Pt(ls)
        if align == "center":
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif align == "right":
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.paragraph_format.space_after = Pt(after)
        p.paragraph_format.space_before = Pt(0)
        if text:
            sf(p.add_run(text), size, bold)
        return p

    def border(cell):
        """セルは罫線を持たず上下中央揃えにする（罫線は表全体で指定）"""
        tcPr = cell._element.get_or_add_tcPr()
        va = OxmlElement("w:vAlign")
        va.set(qn("w:val"), "center")
        tcPr.append(va)

    def table_borders(t):
        tblPr = t._tbl.tblPr
        # セルの上下余白 1mm
        mar = OxmlElement("w:tblCellMar")
        for e, w in (("top", 57), ("bottom", 57), ("left", 108), ("right", 108)):
            el = OxmlElement(f"w:{e}")
            el.set(qn("w:w"), str(w))
            el.set(qn("w:type"), "dxa")
            mar.append(el)
        tblPr.append(mar)
        bd = OxmlElement("w:tblBorders")
        for e in ("top", "left", "bottom", "right", "insideH", "insideV"):
            el = OxmlElement(f"w:{e}")
            el.set(qn("w:val"), "single")
            el.set(qn("w:sz"), "6")
            el.set(qn("w:color"), "555555")
            bd.append(el)
        tblPr.append(bd)

    def shade(cell):
        tcPr = cell._element.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:fill"), "EFEFEF")
        tcPr.append(shd)

    def table(doc, rows, wide):
        # body の末尾は sectPr なので、その手前の要素を見る
        elems = [e for e in doc.element.body if not e.tag.endswith("}sectPr")]
        last = elems[-1] if elems else None
        if (
            last is not None
            and last.tag.endswith("}p")
            and not (last.xpath("string(.)") == " ")
        ):
            gap = doc.add_paragraph()
            gap.paragraph_format.space_after = Pt(0)
            gap.paragraph_format.line_spacing = Pt(5)
            sf(gap.add_run(" "), 1)
        t = doc.add_table(rows=0, cols=2)
        t.alignment = WD_TABLE_ALIGNMENT.CENTER
        t.autofit = False
        lw = 52 if wide else 42
        # 列幅は gridCol と tblLayout=fixed に書かないと LibreOffice/Word で半々に割られる
        t.columns[0].width = Mm(lw)
        t.columns[1].width = Mm(174 - lw)
        lay = OxmlElement("w:tblLayout")
        lay.set(qn("w:type"), "fixed")
        t._tbl.tblPr.append(lay)
        table_borders(t)
        for r in rows:
            lab, val = (r["label"], r["value"]) if isinstance(r, dict) else r
            hm = re.search(r"\{h(\d+)\}", val)
            val = re.sub(r"\{h\d+\}", "", val)
            row = t.add_row()
            row.height = Mm(int(hm.group(1)) if hm else (13 if val == "" else 8))
            c0, c1 = row.cells
            c0.width, c1.width = Mm(lw), Mm(174 - lw)
            for c in (c0, c1):
                border(c)
                c.paragraphs[0].paragraph_format.space_after = Pt(0)
            no_autospace(c0.paragraphs[0])
            c0.paragraphs[0].paragraph_format.line_spacing = Pt(13)
            sf(c0.paragraphs[0].add_run(lab), 9.5, True)
            lines = cell_lines(val, 35 if wide else 38)
            for i, ln in enumerate(lines):
                p = c1.paragraphs[0] if i == 0 else c1.add_paragraph()
                no_autospace(p)
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = Pt(13)
                sf(p.add_run(ln), 9)
        tail = doc.add_paragraph()
        tail.paragraph_format.space_after = Pt(0)
        tail.paragraph_format.line_spacing = Pt(7)
        sf(tail.add_run(" "), 1)

    doc = Document()
    # 文書のプロパティ（python-docx の雛形は作成者 python-docx・2013年のままなので上書きする）
    cp = doc.core_properties
    cp.title = f["title"]
    cp.author = META.get("issuer", "")
    cp.last_modified_by = META.get("issuer", "")
    cp.comments = ""
    cp.created = _dt.datetime.fromisoformat(META["as_of"])
    cp.modified = cp.created
    sec = doc.sections[0]
    sec.page_width, sec.page_height = Mm(210), Mm(297)
    sec.left_margin = sec.right_margin = Mm(18)
    sec.top_margin, sec.bottom_margin = Mm(15), Mm(13)
    st = doc.styles["Normal"]
    st.font.name = FONT
    st.font.size = Pt(10)
    st.element.rPr.rFonts.set(qn("w:eastAsia"), FONT)

    para(doc, f["title"], 16, True, "center", 6, ls=22)
    para(doc, DATE_W, 10, align="right", after=4)
    to = f["to"]
    if "addr" in f:
        for i, x in enumerate(f["addr"]):
            para(doc, flat(x), after=7 if i == len(f["addr"]) - 1 else 0)
    elif to in ("会社", "保証人"):
        para(doc, CO, after=0)
        para(doc, "代表取締役　" + "＿" * 10 + "　殿", after=7)
    elif to == "本人":
        para(doc, "＿" * 14 + "　殿", after=2)
        para(doc, CO, align="right", after=0)
        para(doc, "所在地：" + "＿" * 17, align="right", after=0)
        para(
            doc, "担当者：" + "＿" * 10 + "　電話：" + "＿" * 8, align="right", after=7
        )
    elif to == "主治医":
        para(doc, "＿" * 12 + "　病院・医院", after=0)
        para(doc, "＿" * 12 + "　先生　御机下", after=7)
    elif to == "会社←主治医":
        para(doc, CO, after=0)
        para(doc, "人事・労務担当者　殿", after=7)
    elif to == "社労士":
        para(doc, "みなの社会保険労務士事務所　行", after=7)
    for s in f["intro"]:
        para(doc, flat(s), after=3)
    wide = is_wide(f)
    company = [blk for blk in f["blocks"] if blk.get("company")]
    for blk in f["blocks"]:
        if blk.get("company"):
            continue
        t = blk["type"]
        if t == "p":
            for ln in flat(blk["text"]).split("\n"):
                para(doc, ln, after=4)
        elif t == "fields":
            table(doc, blk["rows"], wide)
        elif t == "checks":
            para(doc, blk["title"], 10, True, after=1)
            for i in blk["items"]:
                p = para(doc, "□　" + i, 9.5, after=1, ls=14)
                p.paragraph_format.left_indent = Mm(4)
        elif t == "box":
            gap = doc.add_paragraph()
            gap.paragraph_format.space_after = Pt(0)
            gap.paragraph_format.line_spacing = Pt(6)
            sf(gap.add_run(" "), 1)
            tb = doc.add_table(rows=2, cols=1)
            tb.alignment = WD_TABLE_ALIGNMENT.CENTER
            table_borders(tb)
            hcell, bcell = tb.rows[0].cells[0], tb.rows[1].cells[0]
            for c in (hcell, bcell):
                border(c)
                c.width = Mm(174)
            hcell.paragraphs[0].paragraph_format.line_spacing = Pt(13)
            hcell.paragraphs[0].paragraph_format.space_after = Pt(0)
            sf(hcell.paragraphs[0].add_run(blk["title"]), 9.5, True)
            tb.rows[1].height = Mm(blk["height"])
            tail = doc.add_paragraph()
            tail.paragraph_format.space_after = Pt(2)
        elif t == "note":
            para(doc, "※ " + blk["text"], 9, after=3, ls=13)
    sig = signature_rows(to, f)
    if sig:
        table(doc, sig, wide)
    for blk in company:
        table(doc, blk["rows"], wide)
    # フッター（右下の社名）は入れない（2026-09-06）
    doc.save(path)


if __name__ == "__main__":
    print(
        "このモジュールは build_shoshiki.py / build_zip.py から import して使う（HTML＝page()、Word＝build_docx()）"
    )
